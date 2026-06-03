import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import os

# === 动态加载自定义字体 ===
# 假设字体文件和 app.py 在同一目录
font_path = "NotoSansCJKsc-Regular.otf"

# 如果你把字体放在 fonts/ 文件夹里，用这行：
# font_path = "fonts/NotoSansCJKsc-Regular.otf"

if os.path.exists(font_path):
    # 添加字体到 Matplotlib 的字体管理器
    font_prop = fm.FontProperties(fname=font_path)
    # 设置全局字体
    plt.rcParams['font.family'] = font_prop.get_name()
    plt.rcParams['axes.unicode_minus'] = False
    st.write("✅ 字体加载成功！")  # 部署后可删除这行
else:
    st.error(f"❌ 字体文件未找到: {font_path}")
    st.stop()  # 停止执行，方便调试

# === 你的其他导入和代码 ===
import numpy as np
import pandas as pd
# ... 其他代码

# 绘图示例
st.title("SPC 控制图")
fig, ax = plt.subplots()
ax.plot([1,2,3], [4,5,6])
ax.set_title("中文标题测试")
ax.set_xlabel("X轴标签")
ax.set_ylabel("Y轴标签")
st.pyplot(fig)
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import StringIO, BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

st.set_page_config(page_title="制药SPC智能分析平台", layout="wide")
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'Microsoft YaHei']
plt.rcParams['axes.unicode_minus'] = False

# ================= 自适应 SPC 计算引擎 (含异常检测) =================
class PharmaSPC:
    E2_IMR, D3_IMR, D4_IMR = 2.660, 0, 3.267
    A2 = {2:1.880, 3:1.023, 4:0.729, 5:0.577, 6:0.483, 7:0.419, 8:0.373}
    D3_R = {2:0, 3:0, 4:0, 5:0, 6:0, 7:0.076, 8:0.136}
    D4_R = {2:3.267, 3:2.574, 4:2.282, 5:2.114, 6:2.004, 7:1.924, 8:1.864}

    def __init__(self, df, value_col, subgroup_col=None):
        self.value_col = value_col
        self.subgroup_col = subgroup_col
        
        # ★ 核心修复：无论是否分组，都保留原始数值序列的引用
        self.data = df[value_col].dropna().reset_index(drop=True)
        
        if subgroup_col and subgroup_col in df.columns:
            grouped = df.groupby(subgroup_col)[value_col].apply(list).reset_index()
            self.subgroups = grouped[value_col].tolist()
            self.labels = grouped[subgroup_col].astype(str).tolist()
            self.n = len(self.subgroups[0]) if self.subgroups else 1
            self.chart_type = "Xbar-S" if self.n > 8 else ("Xbar-R" if self.n >= 2 else "I-MR")
            # 如果子组内只有1个值，强制降级为I-MR
            if self.n < 2:
                self.mr = self.data.diff().abs().dropna().reset_index(drop=True)
                self.chart_type = "I-MR"
        else:
            self.mr = self.data.diff().abs().dropna().reset_index(drop=True)
            self.labels = list(range(1, len(self.data) + 1))
            self.n = 1
            self.chart_type = "I-MR"
            
        self._calc_limits()
        self._detect_violations()

    def _calc_limits(self):
        if self.chart_type == "I-MR":
            self.x_bar = self.data.mean()
            self.mr_bar = self.mr.mean()
            self.ucl_x = self.x_bar + self.E2_IMR * self.mr_bar
            self.lcl_x = self.x_bar - self.E2_IMR * self.mr_bar
            self.ucl_sub = self.D4_IMR * self.mr_bar
            self.sub_label, self.sub_center = "MR", self.mr_bar
        elif self.chart_type == "Xbar-R":
            means = [np.mean(g) for g in self.subgroups]
            ranges = [np.max(g)-np.min(g) for g in self.subgroups]
            self.x_bar = np.mean(means)
            self.r_bar = np.mean(ranges)
            n_key = min(max(self.n, 2), 8)
            self.ucl_x = self.x_bar + self.A2[n_key] * self.r_bar
            self.lcl_x = self.x_bar - self.A2[n_key] * self.r_bar
            self.ucl_sub = self.D4_R[n_key] * self.r_bar
            self.sub_values, self.sub_center, self.sub_label = ranges, self.r_bar, "R"
        else:  # Xbar-S
            means = [np.mean(g) for g in self.subgroups]
            stds = [np.std(g, ddof=1) for g in self.subgroups]
            self.x_bar = np.mean(means)
            self.s_bar = np.mean(stds)
            # 简化常数取值，实际使用建议查完整表
            self.ucl_x = self.x_bar + 1.032 * self.s_bar  
            self.lcl_x = self.x_bar - 1.032 * self.s_bar
            self.ucl_sub = 1.816 * self.s_bar              
            self.sub_values, self.sub_center, self.sub_label = stds, self.s_bar, "S"

    def _detect_violations(self):
        """检测 Western Electric Rule 1: 超出控制限"""
        self.violations = []
        if self.chart_type == "I-MR":
            values = self.data.values
            viol_labels = self.labels
        else:
            values = np.array([np.mean(g) for g in self.subgroups])
            viol_labels = self.labels
            
        for i, val in enumerate(values):
            if val > self.ucl_x or val < self.lcl_x:
                label = viol_labels[i] if i < len(viol_labels) else i+1
                direction = "超出UCL ⬆️" if val > self.ucl_x else "低于LCL ⬇️"
                self.violations.append({
                    "序号/批次": str(label),
                    "测量值/均值": round(float(val), 4),
                    "异常类型": direction,
                    "偏离程度": round(abs(val - self.x_bar), 4)
                })

    @property
    def sample_count(self):
        """统一获取用于报告的样本/子组数量"""
        if self.chart_type == "I-MR":
            return len(self.data)
        return len(self.subgroups)

    def plot(self):
        fig, axes = plt.subplots(2, 1, figsize=(10, 7), sharex=True,
                                  gridspec_kw={'height_ratios': [3, 1]})
        
        if self.chart_type == "I-MR":
            x = np.arange(1, len(self.data)+1)
            vals = self.data.values
            sub_vals = self.mr.values
            sub_x = np.arange(2, len(self.mr)+2)
        else:
            x = np.arange(1, len(self.subgroups)+1)
            vals = np.array([np.mean(g) for g in self.subgroups])
            sub_vals = self.sub_values
            sub_x = x

        # 主图绘制
        normal_mask = ~((vals > self.ucl_x) | (vals < self.lcl_x))
        axes[0].plot(x[normal_mask], vals[normal_mask], 'bo-', markersize=6, label='正常点')
        violation_idx = np.where(~normal_mask)[0]
        if len(violation_idx) > 0:
            axes[0].plot(x[violation_idx], vals[violation_idx], 'ro', markersize=10, 
                         markeredgecolor='darkred', markeredgewidth=2, label=f'异常点(n={len(violation_idx)})')
            for idx in violation_idx:
                axes[0].annotate(f'{vals[idx]:.2f}', (x[idx], vals[idx]), 
                                textcoords="offset points", xytext=(0,12), ha='center',
                                fontsize=9, color='red', fontweight='bold')
                
        axes[0].axhline(self.x_bar, color='green', linestyle='-', lw=1.5, label=f'CL={self.x_bar:.2f}')
        axes[0].axhline(self.ucl_x, color='red', linestyle='--', lw=1.5, label=f'UCL={self.ucl_x:.2f}')
        axes[0].axhline(self.lcl_x, color='red', linestyle='--', lw=1.5, label=f'LCL={self.lcl_x:.2f}')
        axes[0].set_title(f'{self.value_col} - {self.chart_type} 控制图', fontsize=14, fontweight='bold')
        axes[0].set_ylabel('测量值/均值'); axes[0].legend(loc='best', fontsize=9); axes[0].grid(True, alpha=0.3)

        # 子图绘制
        axes[1].plot(sub_x, sub_vals, 'rs-', markersize=5)
        axes[1].axhline(self.sub_center, color='green', linestyle='-', lw=1.5, label=f'{self.sub_label}̄={self.sub_center:.2f}')
        axes[1].axhline(self.ucl_sub, color='red', linestyle='--', lw=1.5, label=f'UCL_{self.sub_label}={self.ucl_sub:.2f}')
        axes[1].set_title(f'{self.sub_label} 图', fontsize=12)
        axes[1].set_xlabel('子组/批次序号'); axes[1].set_ylabel(self.sub_label)
        axes[1].legend(fontsize=9); axes[1].grid(True, alpha=0.3)
        
        plt.tight_layout()
        return fig

# ================= Word 报告生成器 (含异常清单) =================
def generate_word_report(spc_obj, fig, df_preview):
    doc = Document()
    title = doc.add_heading('SPC 统计过程控制分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ★ 使用统一的 sample_count 属性避免 AttributeError
    info = [("分析指标", spc_obj.value_col), ("图表类型", spc_obj.chart_type),
            ("报告时间", datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("样本量/子组数", str(spc_obj.sample_count)),
            ("异常点数量", f"{len(spc_obj.violations)} 个")]
    t = doc.add_table(rows=len(info), cols=2); t.style='Table Grid'
    for i,(k,v) in enumerate(info): 
        t.cell(i,0).text=k; t.cell(i,1).text=v
        for p in t.row_cells(i)[0].paragraphs:
            for r in p.runs: r.bold = True
    
    doc.add_paragraph("")
    doc.add_heading('1. 控制限摘要', level=1)
    doc.add_paragraph(f"• UCL: {spc_obj.ucl_x:.4f}   • CL: {spc_obj.x_bar:.4f}   • LCL: {spc_obj.lcl_x:.4f}")
    
    # 异常数据清单
    doc.add_heading('2. 异常数据识别 (Rule 1: 超出控制限)', level=1)
    if len(spc_obj.violations) == 0:
        p = doc.add_paragraph("✅ 当前数据集未检测到超出控制限的异常点，过程处于统计受控状态。")
        p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
    else:
        warn = doc.add_paragraph(f"⚠️ 共发现 {len(spc_obj.violations)} 个异常点，请调查根本原因：")
        warn.runs[0].font.color.rgb = RGBColor(204, 0, 0)
        vt = doc.add_table(rows=1, cols=4); vt.style='Table Grid'
        headers = ["序号/批次", "测量值/均值", "异常类型", "偏离程度"]
        for i,h in enumerate(headers): 
            vt.rows[0].cells[i].text = h
            for p in vt.rows[0].cells[i].paragraphs:
                for r in p.runs: r.bold = True; r.font.size = Pt(9)
        for v in spc_obj.violations:
            rc = vt.add_row().cells
            for i,k in enumerate(headers): 
                rc[i].text = str(v[k])
                for p in rc[i].paragraphs:
                    for r in p.runs: r.font.size = Pt(9)
    
    doc.add_heading('3. 控制图', level=1)
    buf = BytesIO(); fig.savefig(buf, format='png', dpi=200, bbox_inches='tight'); buf.seek(0)
    doc.add_picture(buf, width=Inches(6.0)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('4. 附录数据', level=1)
    dt = doc.add_table(rows=1, cols=len(df_preview.columns)); dt.style='Table Grid'
    for i,c in enumerate(df_preview.columns): dt.rows[0].cells[i].text=str(c)
    for _,r in df_preview.iterrows():
        rc = dt.add_row().cells
        for i,v in enumerate(r): rc[i].text=str(v)
        
    doc.add_paragraph("")
    disc = doc.add_paragraph("【合规声明】本报告由内部辅助工具自动生成，未经CSV验证，不可作为GMP正式放行依据。异常点判定仅基于Western Electric Rule 1，完整趋势分析需结合其他判异规则。")
    disc.runs[0].font.size = Pt(8)
    
    wb = BytesIO(); doc.save(wb); wb.seek(0)
    return wb

# ================= Web 界面 =================
st.title("📊 制药SPC智能分析平台（异常标注版）")
st.caption("自动识别异常点 | 图表红点高亮 | Word报告含异常清单 | 支持 I-MR / Xbar-R / Xbar-S")

st.subheader("1️⃣ 粘贴您的数据")
raw_text = st.text_area("👇 粘贴数据（含表头）：", height=180,
    placeholder="批号\t含量\nB01\t98.2\nB02\t99.1\nB03\t105.6  ← 异常示例\nB04\t97.8")

df = None
if raw_text.strip():
    try:
        for sep in ['\t', ',', r'\s+']:
            try:
                df = pd.read_csv(StringIO(raw_text), sep=sep, engine='python')
                if df.shape[1] > 1: break
            except: continue
        if df is None or df.shape[1] <= 1:
            st.error("❌ 无法解析，请确认含表头且为表格格式")
        else:
            for col in df.columns[1:]:
                df[col] = pd.to_numeric(df[col].astype(str).str.replace('%','').str.strip(), errors='coerce')
            st.success(f"✅ 解析成功: {df.shape[0]}行 × {df.shape[1]}列")
            with st.expander("🔍 预览数据"): st.dataframe(df, use_container_width=True)
    except Exception as e: st.error(f"❌ 异常: {e}")

if df is not None and df.shape[1] > 1:
    st.subheader("2️⃣ 配置分析参数")
    all_cols = df.columns.tolist()
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    
    c1, c2 = st.columns(2)
    with c1: value_col = st.selectbox("选择分析指标（数值列）", numeric_cols)
    with c2: 
        subgroup_col = st.selectbox("选择子组标识列（可选）", 
                                     options=["(不使用子组 → I-MR图)"] + [c for c in all_cols if c != value_col])
        subgroup_col = None if "(不使用" in subgroup_col else subgroup_col
    
    if st.button("🚀 生成控制图与报告", type="primary"):
        spc = PharmaSPC(df=df, value_col=value_col, subgroup_col=subgroup_col)
        fig = spc.plot()
        st.pyplot(fig)
        
        mc1, mc2, mc3, mc4, mc5 = st.columns(5)
        mc1.metric("图表类型", spc.chart_type)
        mc2.metric("CL", f"{spc.x_bar:.3f}")
        mc3.metric("UCL", f"{spc.ucl_x:.3f}")
        mc4.metric("LCL", f"{spc.lcl_x:.3f}")
        mc5.metric("⚠️ 异常点数", f"{len(spc.violations)}", delta_color="inverse")
        
        # 网页端展示异常清单
        if len(spc.violations) > 0:
            st.warning(f"⚠️ 检测到 {len(spc.violations)} 个异常点（超出控制限）：")
            st.dataframe(pd.DataFrame(spc.violations), use_container_width=True, hide_index=True)
        else:
            st.success("✅ 未检测到超出控制限的异常点")
        
        st.subheader("3️⃣ 导出报告")
        report_buf = generate_word_report(spc, fig, df)
        fname = f"SPC_{spc.chart_type}_{value_col}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        st.download_button("📥 下载 Word 分析报告（含异常标注）", report_buf, fname,
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="secondary")
